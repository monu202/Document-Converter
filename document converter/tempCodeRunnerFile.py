from flask import Flask, request, send_file, render_template
from werkzeug.utils import secure_filename
import os
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF

app = Flask(__name__, template_folder='templates')

# Configuration for upload and converted folders
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['CONVERTED_FOLDER'] = 'converted'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx'}

# Ensure folders exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['CONVERTED_FOLDER'], exist_ok=True)
os.makedirs('fonts', exist_ok=True)  # Ensure fonts folder exists

# Function to check allowed file extensions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/', methods=['GET'])
def index():
    # Render the index page
    return render_template('base.html')

@app.route('/convert', methods=['POST'])
def convert_file():
    # Check if a file is uploaded
    if 'file' not in request.files:
        return "No file part in the request.", 400

    file = request.files['file']
    if file.filename == '':
        return "No file selected.", 400

    # Check for allowed file extensions
    if not allowed_file(file.filename):
        return "Invalid file type selected. Please upload a PDF or Word file.", 400

    conversion_type = request.form['conversionType']
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    # Process the conversion based on the selected type
    if conversion_type == 'pdf-to-word':
        output_path = convert_pdf_to_word(filepath, filename)
    elif conversion_type == 'word-to-pdf':
        output_path = convert_word_to_pdf(filepath, filename)
    else:
        return "Invalid conversion type selected.", 400

    # Send the converted file as a download
    return send_file(output_path, as_attachment=True)

def convert_pdf_to_word(filepath, filename):
    # Convert PDF to Word
    output_path = os.path.join(app.config['CONVERTED_FOLDER'], filename.replace('.pdf', '.docx'))
    doc = Document()
    pdf_reader = PdfReader(filepath)

    for page in pdf_reader.pages:
        doc.add_paragraph(page.extract_text())

    doc.save(output_path)
    return output_path

def convert_word_to_pdf(filepath, filename):
    # Convert Word to PDF
    output_path = os.path.join(app.config['CONVERTED_FOLDER'], filename.replace('.docx', '.pdf'))
    pdf = FPDF()
    pdf.add_page()

    # Add a Unicode font (TrueType font)
    font_path = 'fonts/DejaVuSans.ttf'  # Ensure the font file is in the "fonts" folder
    if not os.path.exists(font_path):
        raise FileNotFoundError(f"Font file not found: {font_path}")
    
    pdf.add_font('DejaVu', '', font_path, uni=True)
    pdf.set_font('DejaVu', size=12)

    # Read the content of the Word file
    doc = Document(filepath)
    for para in doc.paragraphs:
        pdf.multi_cell(0, 10, para.text)

    pdf.output(output_path)
    return output_path

if __name__ == '__main__':
    app.run(debug=True) 